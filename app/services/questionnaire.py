import json
import os
import re
import tempfile
import asyncio
from typing import List, Tuple, Dict

import requests
import aiohttp

from PyPDF2 import PdfReader
from docx import Document

try:
    import textract  # type: ignore
except Exception:  # pragma: no cover - optional dependency may be missing
    textract = None
from fpdf import FPDF

from app.utils.logger import logger
from app.config import settings

# Cache for generated answers to avoid repeated token usage
QA_CACHE: Dict[Tuple[str, str], str] = {}

# Maximum words from context to send to the model
DEFAULT_MAX_CONTEXT_WORDS = 250

# Cheaper model for question answering
CHEAPER_MODEL = "gpt-4o"

def _trim_context(context: str, max_words: int = DEFAULT_MAX_CONTEXT_WORDS) -> str:
    """Return a truncated context limited to ``max_words`` words."""
    words = context.split()
    if len(words) <= max_words:
        return context
    logger.info(f"Trimming context from {len(words)} to {max_words} words")
    return " ".join(words[:max_words])

def _extract_text(file_bytes: bytes, filename: str) -> str:
    ext = os.path.splitext(filename)[1].lower()
    logger.info(f"Extracting text from {filename}")
    if ext == ".pdf":
        with tempfile.NamedTemporaryFile(delete=False, suffix=".pdf") as tmp:
            tmp.write(file_bytes)
            tmp.flush()
            reader = PdfReader(tmp.name)
            text = "\n".join(page.extract_text() or "" for page in reader.pages)
        os.remove(tmp.name)
        logger.info(f"Successfully extracted text from {filename}")
        return text
    elif ext in [".docx", ".doc"]:
        with tempfile.NamedTemporaryFile(delete=False, suffix=ext) as tmp:
            tmp.write(file_bytes)
            tmp.flush()
            try:
                doc = Document(tmp.name)
                text = "\n".join(p.text for p in doc.paragraphs)
            except Exception:
                if textract:
                    try:
                        text = textract.process(tmp.name).decode(
                            "utf-8", errors="ignore"
                        )
                    except Exception as e:  # pragma: no cover - fallback path
                        logger.error(f"Failed to parse {filename}: {e}")
                        text = ""
                else:
                    logger.error(f"doc parsing requires textract: {filename}")
                    text = ""
        os.remove(tmp.name)
        logger.info(f"Successfully extracted text from {filename}")
        return text
    elif ext in [".txt", ".md"]:
        text = file_bytes.decode("utf-8", errors="ignore")
        logger.info(f"Successfully extracted text from {filename}")
        return text
    else:
        logger.warning(f"Unsupported file type for {filename}")
        raise ValueError("Unsupported file type")


def _extract_questions(text: str) -> List[str]:
    """Parse questions from raw text.

    The previous implementation relied heavily on newlines which caused
    incomplete extraction for some PDFs/Docs where each line was split in an
    odd way (often leaving only ``?``).  To make the extraction more robust we
    normalise whitespace and use a simple regular expression to capture any
    segment that ends with a question mark.
    """

    # Collapse all whitespace to single spaces so regex can span lines/bullets
    cleaned = re.sub(r"\s+", " ", text)

    # Find sequences ending with a question mark.  ``re.findall`` will return
    # all non-empty matches so even short questions like "Why?" are preserved.
    raw_questions = re.findall(r"[^?]+?\?", cleaned)

    questions = [q.strip() for q in raw_questions if q.strip().endswith("?")]

    logger.info(f"Parsed {len(questions)} questions")
    return questions



def _split_paragraphs(text: str) -> List[str]:
    paras = [p.strip() for p in re.split(r"\n{2,}", text) if p.strip()]
    return paras


def _search_documents(
    question: str, documents: List[Tuple[str, str]]
) -> Tuple[str, str]:
    q_words = set(re.findall(r"\w+", question.lower()))
    best_score = 0
    best_para = ""
    best_file = ""
    for filename, content in documents:
        for para in _split_paragraphs(content):
            words = set(re.findall(r"\w+", para.lower()))
            score = len(q_words & words)
            if score > best_score:
                best_score = score
                best_para = para
                best_file = filename
    return best_para, best_file


def process_questionnaire(
    file_bytes: bytes, filename: str, documents: List[Tuple[str, str]]
):
    logger.info(f"Processing questionnaire for file: {filename}")
    text = _extract_text(file_bytes, filename)
    questions = _extract_questions(text)
    results = []
    for q in questions:
        para, doc_name = _search_documents(q, documents)
        answer = para if para else "No matching information found."
        reference = f"({doc_name}, snippet: {para[:20]}...)" if para else ""
        results.append(
            {
                "question": q,
                "answer": answer,
                "source_file": doc_name,
                "source_paragraph": para[:100] + "..." if len(para) > 100 else para,
                "reference": reference,
            }
        )
    logger.info(f"Processed {len(results)} questions for file: {filename}")
    return results


async def _query_vector_store(
    question: str, vector_store_id: str, file_map: Dict[str, str], openai_api_key: str
) -> Tuple[str, str, str, int]:
    """
    Query the vector store using an HTTP call to retrieve relevant chunks.
    
    Args:
        question (str): The question to query the vector store with.
        vector_store_id (str): The ID of the vector store.
        file_map (Dict[str, str]): A mapping of file IDs to filenames.
        openai_api_key (str): The OpenAI API key for authentication.
    
    Returns:
        Tuple[str, str, str, int]: Chunk text, filename, snippet, and page number.
    """
    logger.info(f"Querying vector store for question: '{question}' with vector_store_id: {vector_store_id}")
    
    # Define the API endpoint for searching
    url = f"https://api.openai.com/v1/vector_stores/{vector_store_id}/search"
    
    # Set up headers with authentication
    headers = {
        "Authorization": f"Bearer {openai_api_key}",
        "Content-Type": "application/json"
    }
    
    # Prepare the request body
    data = {
        "query": question,
    }
    
    try:
        # Send the POST request
        response = requests.post(url, headers=headers, data=json.dumps(data))
        
        # Check for successful response
        if response.status_code == 200:
            resp_data = response.json()
            if resp_data.get("data") and len(resp_data["data"]) > 0:
                record = resp_data["data"][0]
                content = record["content"][0]
                chunk = content.get("text", "")
                meta = content.get("metadata", {})
                page_no = int(meta.get("page", 1))
                if not chunk:
                    logger.warning(f"Empty chunk received for question '{question}'")
                file_id = record.get("file_id", "")
                filename = file_map.get(file_id, file_id)
                logger.info(
                    f"Found match for question '{question}': file_id {file_id}, filename {filename}, page {page_no}"
                )
                return chunk, filename, chunk[:20], page_no
            else:
                logger.info(f"No match found in vector store for question '{question}'")
                return "", "", "", 0
        else:
            logger.error(f"API request failed with status code {response.status_code}: {response.text}")
            return "", "", "", 0
    
    except Exception as e:
        logger.error(f"Vector store query failed for question '{question}': {e}")
        return "", "", "", 0
        
async def _generate_answer(question: str, context: str, openai_client) -> str:
    logger.info(f"Generating answer for question: '{question}'")
    if not context:
        logger.warning(f"Cannot generate answer for '{question}' due to empty context.")
        return ""

    cached = QA_CACHE.get((question, context))
    if cached:
        logger.info(f"Using cached answer for question '{question}'")
        return cached

    short_context = _trim_context(context)

    try:
        resp = await asyncio.to_thread(
            openai_client.chat.completions.create,
            model=CHEAPER_MODEL,
            messages=[
                {
                    "role": "system",
                    "content": "Answer the user question using the provided context as reference.",
                },
                {
                    "role": "user",
                    "content": f"Context:\n{short_context}\n\nQuestion: {question}",
                },
            ],
        )
        if resp.choices:
            answer = resp.choices[0].message.content.strip()
            QA_CACHE[(question, context)] = answer
            tokens = getattr(getattr(resp, "usage", None), "total_tokens", "unknown")
            logger.info(
                f"Successfully generated answer for question '{question}' using {tokens} tokens"
            )
            return answer
    except Exception as e:  # pragma: no cover - external call
        logger.error(f"Failed to generate answer for '{question}': {e}")
    return ""


async def process_questionnaire_rag(
    file_bytes: bytes,
    filename: str,
    vector_store_id: str,
    file_map: Dict[str, str],
    openai_client,
) -> List[dict]:
    logger.info(f"Processing questionnaire_rag for file: {filename} with vector_store_id: {vector_store_id}")
    text = _extract_text(file_bytes, filename)
    questions = _extract_questions(text)
    results = []
    for q in questions:
        para, fname, snippet, page_no = await _query_vector_store(
            q, vector_store_id, file_map, settings.OPENAI_API_KEY
        )
        answer = "No matching information found."
        if para:
            gen = await _generate_answer(q, para, openai_client)
            answer = gen or para
        reference = f"({fname}, {page_no}, {snippet}...)" if para else ""
        results.append(
            {
                "question": q,
                "answer": answer,
                "source_file": fname,
                "source_filename": fname,
                "source_paragraph": para[:100] + "..." if len(para) > 100 else para,
                "reference": reference,
            }
        )
    logger.info(f"Processed {len(results)} questions (RAG) for file: {filename}")
    return results


def generate_docx(results: List[dict]) -> bytes:
    logger.info(f"Generating DOCX for {len(results)} results.")
    from docx import Document as DocxDoc

    doc = DocxDoc()
    try:
        for item in results:
            doc.add_paragraph(item["question"]).bold = True
            doc.add_paragraph(item["answer"])
            if item.get("reference"):
                doc.add_paragraph(f"Source: {item['source_file']} ")
                doc.add_paragraph(f"Paragraph: {item['source_paragraph'][:200]}...")
            doc.add_paragraph("")
        buf = tempfile.NamedTemporaryFile(delete=False, suffix=".docx")
        doc.save(buf.name)
        with open(buf.name, "rb") as f:
            data = f.read()
        os.remove(buf.name)
        logger.info("Successfully generated DOCX.")
        return data
    except Exception as e:
        logger.error(f"Error generating DOCX: {e}")
        raise


def generate_pdf(results: List[dict]) -> bytes:
    logger.info(f"Generating PDF for {len(results)} results.")
    pdf = FPDF()
    try:
        pdf.add_page()
        pdf.set_font("Arial", size=12)
        for item in results:
            pdf.set_font(style="B", family="Arial", size=12)
            pdf.multi_cell(0, 10, item["question"])
            pdf.set_font(style="", family="Arial", size=12)
            pdf.multi_cell(0, 10, item["answer"])
            if item.get("reference"):
                pdf.set_font(style="I", family="Arial", size=11)
                pdf.multi_cell(0, 10, f"Source: {item['reference']}")
            pdf.ln()
        data = pdf.output(dest="S").encode("latin-1")
        logger.info("Successfully generated PDF.")
        return data
    except Exception as e:
        logger.error(f"Error generating PDF: {e}")
        raise